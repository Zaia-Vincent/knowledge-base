"""Multi-format text extractor — extracts text from PDF, DOCX, XLSX, MSG, images, and plain text."""

import logging
import mimetypes
from pathlib import Path

logger = logging.getLogger(__name__)


class MultiFormatTextExtractor:
    """Infrastructure adapter that extracts text from various file formats.

    Implements the TextExtractor interface using format-specific libraries:
    - PDF: PyMuPDF (fitz)
    - DOCX: python-docx
    - XLSX: openpyxl
    - MSG: extract-msg
    - TXT/CSV/MD: built-in
    - Images: placeholder (will use LLM vision in Sprint 3+)
    """

    # Format → handler method mapping
    _HANDLERS: dict[str, str] = {
        "application/pdf": "_extract_pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_extract_docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "_extract_xlsx",
        "application/vnd.ms-outlook": "_extract_msg",
        "text/plain": "_extract_text",
        "text/csv": "_extract_text",
        "text/markdown": "_extract_text",
        "text/html": "_extract_text",
        "image/png": "_extract_image",
        "image/jpeg": "_extract_image",
        "image/tiff": "_extract_image",
        "image/webp": "_extract_image",
    }

    def can_extract(self, mime_type: str) -> bool:
        """Check if this extractor supports the given MIME type."""
        return mime_type in self._HANDLERS

    def supported_types(self) -> list[str]:
        """Return all supported MIME types."""
        return list(self._HANDLERS.keys())

    async def extract_text(self, file_path: str, mime_type: str | None = None) -> str:
        """Extract text from a file at the given path.

        Args:
            file_path: Absolute path to the file.
            mime_type: Optional MIME type override. If not given, detected from extension.

        Returns:
            Extracted text content.

        Raises:
            ValueError: If the file type is not supported.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if mime_type is None:
            mime_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"

        handler_name = self._HANDLERS.get(mime_type)
        if handler_name is None:
            raise ValueError(f"Unsupported file type: {mime_type} ({path.name})")

        handler = getattr(self, handler_name)
        text = await handler(file_path)

        logger.info(
            "Extracted %d characters from %s (%s)",
            len(text),
            path.name,
            mime_type,
        )
        return text

    # ── Format-specific handlers ─────────────────────────────────────

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF."""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        pages: list[str] = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(text)
            else:
                # Page might be scanned/image → flag for LLM OCR later
                logger.debug("Page %d appears to be scanned (no text layer)", page_num + 1)

        doc.close()

        if not pages:
            logger.warning("PDF has no extractable text — may require OCR: %s", file_path)
            return ""

        return "\n\n".join(pages)

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    async def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX using openpyxl."""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")

            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    parts.append(" | ".join(cells))

        wb.close()
        return "\n".join(parts)

    async def _extract_msg(self, file_path: str) -> str:
        """Extract text from Outlook MSG files."""
        import extract_msg

        msg = extract_msg.Message(file_path)
        parts = [
            f"From: {msg.sender or ''}",
            f"To: {msg.to or ''}",
            f"Subject: {msg.subject or ''}",
            f"Date: {msg.date or ''}",
            "",
            msg.body or "",
        ]
        msg.close()
        return "\n".join(parts)

    async def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text files (TXT, CSV, MD, HTML)."""
        path = Path(file_path)
        # Try UTF-8 first, then fall back to latin-1
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    async def _extract_image(self, file_path: str) -> str:
        """Placeholder for image OCR — will use LLM vision in a later sprint."""
        logger.info("Image detected: %s — OCR via LLM vision (not yet implemented)", file_path)
        return "[Image file — OCR pending LLM vision integration]"
